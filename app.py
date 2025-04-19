# import streamlit as st
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# from langchain.text_splitter import CharacterTextSplitter
# import os
# from langchain.embeddings import HuggingFaceEmbeddings  
# from langchain.vectorstores import FAISS
# import google.generativeai as genai
# from htmlTemplate import css, bot_template, user_template


# load_dotenv()

# genai.configure(api_key=os.getenv("AIzaSyBhZsqcV6xWWMJg2_lszGkLUGxYQllMv2g")) 



# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             text += page.extract_text() or "" 
#     return text

# def get_text_chunks(raw_text):
#     text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=200, length_function=len)
#     return text_splitter.split_text(raw_text)

# def get_vectorstore(text_chunks):
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")  
#     return FAISS.from_texts(texts=text_chunks, embedding=embeddings)

# def chat_with_gemini(question, chat_history):
#     model = genai.GenerativeModel("gemini-1.5-pro")
#     history = "\n".join([f"User: {msg['content']}" if msg['role'] == 'user' else f"Bot: {msg['content']}" for msg in chat_history])
#     prompt = f"{history}\nUser: {question}\nBot:"  
#     response = model.generate_content(prompt)
#     return response.text

# def handle_userinput(user_question):
#     if st.session_state.vectorstore is not None:
#         retriever = st.session_state.vectorstore.as_retriever()
#         docs = retriever.get_relevant_documents(user_question)
#         context = "\n".join([doc.page_content for doc in docs])
#         full_prompt = f"Context: {context}\n\nQuestion: {user_question}"
#         response = chat_with_gemini(full_prompt, st.session_state.chat_history)

#         st.session_state.chat_history.append({"role": "user", "content": user_question})
#         st.session_state.chat_history.append({"role": "bot", "content": response})

#         for msg in st.session_state.chat_history:
#             if msg['role'] == 'user':
#                 st.write(user_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#             else:
#                 st.write(bot_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#     else:
#         st.warning("Please upload and process a PDF first.")

# def main():
#     st.set_page_config(page_title="Chat with MultiPDFs", page_icon="📚")
#     st.write(css, unsafe_allow_html=True)

#     if "vectorstore" not in st.session_state:
#         st.session_state.vectorstore = None

#     if "chat_history" not in st.session_state:
#         st.session_state.chat_history = []

#     st.header("Chat with MultiPDFs 📚")
#     user_question = st.text_input("Ask your question about your document")

#     if user_question:
#         handle_userinput(user_question)

#     with st.sidebar:
#         st.subheader("Your documents")
#         pdf_docs = st.file_uploader("Upload your PDFs here and click on Process", accept_multiple_files=True)

#         if st.button("Process"):
#             if pdf_docs:
#                 with st.spinner("Processing..."):
#                     raw_text = get_pdf_text(pdf_docs)
#                     text_chunks = get_text_chunks(raw_text)
#                     vectorstore = get_vectorstore(text_chunks)
#                     st.session_state.vectorstore = vectorstore
#                     st.success("Processing complete! You can now chat with your PDFs.")
#             else:
#                 st.warning("Please upload at least one PDF.")

# if __name__ == "__main__":
#     main()




# import streamlit as st
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# from langchain.text_splitter import CharacterTextSplitter
# import os
# from langchain.embeddings import HuggingFaceEmbeddings  
# from langchain.vectorstores import FAISS
# import google.generativeai as genai
# from htmlTemplate import css, bot_template, user_template

# # Load environment variables
# load_dotenv()

# # Configure API key for Google Generative AI
# genai.configure(api_key=os.getenv("AIzaSyCJMs3PeOwVub-K1402Tx_vHqFyijbaILE"))

# # Function to extract text from uploaded PDF documents
# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             text += page.extract_text() or "" 
#     return text

# # Function to split text into chunks
# def get_text_chunks(raw_text):
#     text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=200, length_function=len)
#     return text_splitter.split_text(raw_text)

# # Function to convert text chunks into a vector store
# def get_vectorstore(text_chunks):
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")  
#     return FAISS.from_texts(texts=text_chunks, embedding=embeddings)

# # Function to handle interaction with Gemini AI
# def chat_with_gemini(question, chat_history):
#     model = genai.GenerativeModel("gemini-1.5-pro")
#     history = "\n".join([f"User: {msg['content']}" if msg['role'] == 'user' else f"Bot: {msg['content']}" for msg in chat_history])
#     prompt = f"{history}\nUser: {question}\nBot:"  
#     response = model.generate_content(prompt)
#     return response.text

# # Function to handle user input and interact with the vector store
# def handle_userinput(user_question):
#     if st.session_state.vectorstore is not None:
#         retriever = st.session_state.vectorstore.as_retriever()
#         docs = retriever.get_relevant_documents(user_question)
#         context = "\n".join([doc.page_content for doc in docs])
#         full_prompt = f"Context: {context}\n\nQuestion: {user_question}"
#         response = chat_with_gemini(full_prompt, st.session_state.chat_history)

#         # Store chat history
#         st.session_state.chat_history.append({"role": "user", "content": user_question})
#         st.session_state.chat_history.append({"role": "bot", "content": response})

#         # Display chat history in the app
#         for msg in st.session_state.chat_history:
#             if msg['role'] == 'user':
#                 st.write(user_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#             else:
#                 st.write(bot_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#     else:
#         st.warning("Please upload and process a PDF first.")

# # Main function to control the app flow
# def main():
#     st.set_page_config(page_title="Chat with MultiPDFs", page_icon="📚")
#     st.write(css, unsafe_allow_html=True)

#     # Initialize session state for vectorstore and chat history
#     if "vectorstore" not in st.session_state:
#         st.session_state.vectorstore = None

#     if "chat_history" not in st.session_state:
#         st.session_state.chat_history = []

#     # Header for the app
#     st.header("Chat with MultiPDFs 📚")
#     user_question = st.text_input("Ask your question about your document")

#     # Handle user question input
#     if user_question:
#         handle_userinput(user_question)

#     # Sidebar for uploading PDFs
#     with st.sidebar:
#         st.subheader("Your documents")
#         pdf_docs = st.file_uploader("Upload your PDFs here and click on Process", accept_multiple_files=True)

#         # Button to process the uploaded PDFs
#         if st.button("Process"):
#             if pdf_docs:
#                 with st.spinner("Processing..."):
#                     raw_text = get_pdf_text(pdf_docs)
#                     text_chunks = get_text_chunks(raw_text)
#                     vectorstore = get_vectorstore(text_chunks)
#                     st.session_state.vectorstore = vectorstore
#                     st.success("Processing complete! You can now chat with your PDFs.")
#             else:
#                 st.warning("Please upload at least one PDF.")

# if __name__ == "__main__":
#     main()


# import streamlit as st
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# from langchain.text_splitter import CharacterTextSplitter
# import os
# from langchain.embeddings import HuggingFaceEmbeddings  
# from langchain.vectorstores import FAISS
# import google.generativeai as genai
# from htmlTemplate import css, bot_template, user_template

# # Load environment variables from .env
# load_dotenv()

# # Configure API key for Google Generative AI
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("API key not found. Make sure your .env file has GOOGLE_API_KEY set.")
#     st.stop()
# genai.configure(api_key=api_key)

# # Extract text from uploaded PDF documents
# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             text += page.extract_text() or "" 
#     return text

# # Split raw text into manageable chunks
# def get_text_chunks(raw_text):
#     text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=200, length_function=len)
#     return text_splitter.split_text(raw_text)

# # Convert chunks into a searchable vector store
# def get_vectorstore(text_chunks):
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")  
#     return FAISS.from_texts(texts=text_chunks, embedding=embeddings)

# # Interact with Gemini model
# def chat_with_gemini(question, chat_history):
#     model = genai.GenerativeModel("gemini-1.5-pro")
#     history = "\n".join([f"User: {msg['content']}" if msg['role'] == 'user' else f"Bot: {msg['content']}" for msg in chat_history])
#     prompt = f"{history}\nUser: {question}\nBot:"  
#     response = model.generate_content(prompt)
#     return response.text

# # Handle user's question
# def handle_userinput(user_question):
#     if st.session_state.vectorstore is not None:
#         retriever = st.session_state.vectorstore.as_retriever()
#         docs = retriever.get_relevant_documents(user_question)
#         context = "\n".join([doc.page_content for doc in docs])
#         full_prompt = f"Context: {context}\n\nQuestion: {user_question}"
#         response = chat_with_gemini(full_prompt, st.session_state.chat_history)

#         # Save to chat history
#         st.session_state.chat_history.append({"role": "user", "content": user_question})
#         st.session_state.chat_history.append({"role": "bot", "content": response})

#         # Display chat
#         for msg in st.session_state.chat_history:
#             if msg['role'] == 'user':
#                 st.write(user_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#             else:
#                 st.write(bot_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#     else:
#         st.warning("Please upload and process a PDF first.")

# # Main Streamlit app
# def main():
#     st.set_page_config(page_title="Chat with MultiPDFs", page_icon="📚")
#     st.write(css, unsafe_allow_html=True)

#     # Session state initialization
#     if "vectorstore" not in st.session_state:
#         st.session_state.vectorstore = None
#     if "chat_history" not in st.session_state:
#         st.session_state.chat_history = []

#     st.header("Chat with MultiPDFs 📚")
#     user_question = st.text_input("Ask your question about your document")

#     if user_question:
#         handle_userinput(user_question)

#     # Upload PDFs in the sidebar
#     with st.sidebar:
#         st.subheader("Your documents")
#         pdf_docs = st.file_uploader("Upload your PDFs here and click on Process", accept_multiple_files=True)

#         if st.button("Process"):
#             if pdf_docs:
#                 with st.spinner("Processing..."):
#                     raw_text = get_pdf_text(pdf_docs)
#                     text_chunks = get_text_chunks(raw_text)
#                     vectorstore = get_vectorstore(text_chunks)
#                     st.session_state.vectorstore = vectorstore
#                     st.success("Processing complete! You can now chat with your PDFs.")
#             else:
#                 st.warning("Please upload at least one PDF.")

# if __name__ == "__main__":
#     main()


# import streamlit as st
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# from langchain.text_splitter import CharacterTextSplitter
# import os
# from langchain.embeddings import HuggingFaceEmbeddings  
# from langchain.vectorstores import FAISS
# import google.generativeai as genai
# from docx import Document
# import pandas as pd
# import csv
# from io import StringIO
# from htmlTemplate import css, bot_template, user_template

# # Load environment variables from .env
# load_dotenv()

# # Configure API key for Google Generative AI
# api_key = os.getenv("GOOGLE_API_KEY")
# if not api_key:
#     st.error("API key not found. Make sure your .env file has GOOGLE_API_KEY set.")
#     st.stop()
# genai.configure(api_key=api_key)

# # Extract text from uploaded PDF documents
# def get_pdf_text(pdf_docs):
#     text = ""
#     for pdf in pdf_docs:
#         pdf_reader = PdfReader(pdf)
#         for page in pdf_reader.pages:
#             text += page.extract_text() or "" 
#     return text

# # Extract text from DOCX documents
# def get_docx_text(docx_docs):
#     text = ""
#     for doc in docx_docs:
#         docx_reader = Document(doc)
#         for para in docx_reader.paragraphs:
#             text += para.text + "\n"
#     return text

# # Extract text from TXT documents
# def get_txt_text(txt_docs):
#     text = ""
#     for txt in txt_docs:
#         text += txt.getvalue().decode("utf-8") + "\n"
#     return text

# # Extract text from CSV documents
# def get_csv_text(csv_docs):
#     text = ""
#     for csv_file in csv_docs:
#         csv_data = pd.read_csv(csv_file)
#         text += csv_data.to_string(index=False) + "\n"
#     return text

# # Split raw text into manageable chunks
# def get_text_chunks(raw_text):
#     text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=200, length_function=len)
#     return text_splitter.split_text(raw_text)

# # Convert chunks into a searchable vector store
# def get_vectorstore(text_chunks):
#     embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")  
#     return FAISS.from_texts(texts=text_chunks, embedding=embeddings)

# # Interact with Gemini model
# def chat_with_gemini(question, chat_history):
#     model = genai.GenerativeModel("gemini-1.5-pro")
#     history = "\n".join([f"User: {msg['content']}" if msg['role'] == 'user' else f"Bot: {msg['content']}" for msg in chat_history])
#     prompt = f"{history}\nUser: {question}\nBot:"  
#     response = model.generate_content(prompt)
#     return response.text

# # Handle user's question
# def handle_userinput(user_question):
#     if st.session_state.vectorstore is not None:
#         retriever = st.session_state.vectorstore.as_retriever()
#         docs = retriever.get_relevant_documents(user_question)
#         context = "\n".join([doc.page_content for doc in docs])
#         full_prompt = f"Context: {context}\n\nQuestion: {user_question}"
#         response = chat_with_gemini(full_prompt, st.session_state.chat_history)

#         # Save to chat history
#         st.session_state.chat_history.append({"role": "user", "content": user_question})
#         st.session_state.chat_history.append({"role": "bot", "content": response})

#         # Display chat with attractive formatting
#         for msg in st.session_state.chat_history:
#             if msg['role'] == 'user':
#                 st.write(user_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#             else:
#                 st.write(bot_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
#     else:
#         st.warning("Please upload and process a document first.")

# # Main Streamlit app
# def main():
#     st.set_page_config(page_title="Chat with MultiDocs", page_icon="📚")
#     st.write(css, unsafe_allow_html=True)

#     # Session state initialization
#     if "vectorstore" not in st.session_state:
#         st.session_state.vectorstore = None
#     if "chat_history" not in st.session_state:
#         st.session_state.chat_history = []

#     st.header("Chat with MultiDocs 📚")
#     user_question = st.text_input("Ask your question about your document")

#     if user_question:
#         handle_userinput(user_question)

#     # Upload files in the sidebar
#     with st.sidebar:
#         st.subheader("Your documents")
#         pdf_docs = st.file_uploader("Upload PDF files here", accept_multiple_files=True, type=["pdf"])
#         docx_docs = st.file_uploader("Upload DOCX files here", accept_multiple_files=True, type=["docx"])
#         txt_docs = st.file_uploader("Upload TXT files here", accept_multiple_files=True, type=["txt"])
#         csv_docs = st.file_uploader("Upload CSV files here", accept_multiple_files=True, type=["csv"])
#         xlsx_docs = st.file_uploader("Upload XLSX files here", accept_multiple_files=True, type=["xlsx"])

#         if st.button("Process"):
#             if pdf_docs or docx_docs or txt_docs or csv_docs or xlsx_docs:
#                 with st.spinner("Processing..."):
#                     raw_text = ""
#                     if pdf_docs:
#                         raw_text += get_pdf_text(pdf_docs)
#                     if docx_docs:
#                         raw_text += get_docx_text(docx_docs)
#                     if txt_docs:
#                         raw_text += get_txt_text(txt_docs)
#                     if csv_docs:
#                         raw_text += get_csv_text(csv_docs)
#                     if xlsx_docs:
#                         for xlsx in xlsx_docs:
#                             xlsx_data = pd.read_excel(xlsx)
#                             raw_text += xlsx_data.to_string(index=False) + "\n"
#                     text_chunks = get_text_chunks(raw_text)
#                     vectorstore = get_vectorstore(text_chunks)
#                     st.session_state.vectorstore = vectorstore
#                     st.success("Processing complete! You can now chat with your documents.")
#             else:
#                 st.warning("Please upload at least one document.")

# if __name__ == "__main__":
#     main()


import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
import os
from langchain.embeddings import HuggingFaceEmbeddings  
from langchain.vectorstores import FAISS
import google.generativeai as genai
from docx import Document
import pandas as pd
from htmlTemplate import css, bot_template, user_template

# Load environment variables from .env
load_dotenv()

# Configure API key for Google Generative AI
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    st.error("API key not found. Make sure your .env file has GOOGLE_API_KEY set.")
    st.stop()
genai.configure(api_key=api_key)

# Extract text from uploaded files based on their type
def extract_text_from_file(uploaded_file):
    text = ""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_type == 'pdf':
            pdf_reader = PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
        elif file_type == 'docx':
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_type == 'txt':
            text = uploaded_file.getvalue().decode("utf-8")
        elif file_type == 'csv':
            df = pd.read_csv(uploaded_file)
            text = df.to_string(index=False)
        elif file_type in ['xls', 'xlsx']:
            df = pd.read_excel(uploaded_file)
            text = df.to_string(index=False)
        else:
            st.warning(f"Unsupported file type: {file_type}")
    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
    
    return text

# Get combined text from all uploaded files
def get_combined_text(uploaded_files):
    combined_text = ""
    for file in uploaded_files:
        combined_text += extract_text_from_file(file) + "\n\n"
    return combined_text

# Split raw text into manageable chunks
def get_text_chunks(raw_text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    return text_splitter.split_text(raw_text)

# Convert chunks into a searchable vector store
def get_vectorstore(text_chunks):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")  
    return FAISS.from_texts(texts=text_chunks, embedding=embeddings)

# Interact with Gemini model
def chat_with_gemini(question, chat_history):
    model = genai.GenerativeModel("gemini-1.5-pro")
    history = "\n".join([
        f"User: {msg['content']}" if msg['role'] == 'user' 
        else f"Bot: {msg['content']}" 
        for msg in chat_history
    ])
    prompt = f"{history}\nUser: {question}\nBot:"  
    response = model.generate_content(prompt)
    return response.text

# Handle user's question
def handle_userinput(user_question):
    if st.session_state.vectorstore is not None:
        retriever = st.session_state.vectorstore.as_retriever()
        docs = retriever.get_relevant_documents(user_question)
        context = "\n".join([doc.page_content for doc in docs])
        full_prompt = f"Context: {context}\n\nQuestion: {user_question}"
        response = chat_with_gemini(full_prompt, st.session_state.chat_history)

        # Save to chat history
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        st.session_state.chat_history.append({"role": "bot", "content": response})

        # Display chat with attractive formatting
        for msg in st.session_state.chat_history:
            if msg['role'] == 'user':
                st.write(user_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
            else:
                st.write(bot_template.replace("{{MSG}}", msg['content']), unsafe_allow_html=True)
    else:
        st.warning("Please upload and process documents first.")

# Main Streamlit app
def main():
    st.set_page_config(page_title="Chat with MultiDocs", page_icon="📚")
    st.write(css, unsafe_allow_html=True)

    # Session state initialization
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    st.header("Chat with MultiDocs 📚")
    user_question = st.text_input("Ask your question about your documents")

    if user_question:
        handle_userinput(user_question)

    # Single file uploader for all document types
    with st.sidebar:
        st.subheader("Your documents")
        uploaded_files = st.file_uploader(
            "Upload your files here", 
            accept_multiple_files=True,
            type=["pdf", "docx", "txt", "csv", "xlsx", "xls"]
        )
        
        if st.button("Process Documents"):
            if uploaded_files:
                with st.spinner("Processing documents..."):
                    try:
                        # Get combined text from all files
                        raw_text = get_combined_text(uploaded_files)
                        
                        # Split into chunks
                        text_chunks = get_text_chunks(raw_text)
                        
                        # Create vector store
                        vectorstore = get_vectorstore(text_chunks)
                        st.session_state.vectorstore = vectorstore
                        
                        # Clear previous chat history when new documents are processed
                        st.session_state.chat_history = []
                        
                        st.success("Documents processed successfully! You can now ask questions.")
                    except Exception as e:
                        st.error(f"Error processing documents: {str(e)}")
            else:
                st.warning("Please upload at least one document.")

if __name__ == "__main__":
    main()